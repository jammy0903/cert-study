"""
정보처리기사_실기_출제현황_뜻포함.md → docx 변환기 (B&W 프린트 최적화).

디자인 원칙
 - 컬러 미사용 (흑백 출력 가정, 회색 음영 + 두께·굵기·심볼로만 구분)
 - 섹션별 강제 페이지 나눔 없음 (연속 흐름)
 - 상태 구분: ● 출제됨 / ◐ 약출제 / ○ 미출제 + 굵기/이탤릭
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path('/home/jammy/projects/cert-study/정보처리기사_실기_출제현황_뜻포함.md')
OUT = Path('/home/jammy/projects/cert-study/정보처리기사_실기_출제현황_뜻포함_BW.docx')


# ─── 그레이스케일 팔레트 ──────────────────────────────────
BLACK        = RGBColor(0x00, 0x00, 0x00)
GRAY_DARK    = RGBColor(0x33, 0x33, 0x33)
GRAY_MID     = RGBColor(0x66, 0x66, 0x66)
GRAY_LIGHT   = RGBColor(0xBB, 0xBB, 0xBB)
HEADER_BG    = RGBColor(0x22, 0x22, 0x22)  # 표 헤더: 짙은 회색 (프린트하면 검정에 가까움)
HEADER_TXT   = RGBColor(0xFF, 0xFF, 0xFF)  # 흰 글자
ROW_ALT_BG   = RGBColor(0xF0, 0xF0, 0xF0)  # 짝수행 아주 옅은 회색
EMPHASIS_BG  = RGBColor(0xE4, 0xE4, 0xE4)  # 출제됨 행 배경 (살짝 진함)


# ─── XML 헬퍼 ────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


def set_cell_border(cell, color='888888', sz=4, top=None, bottom=None):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    sides = {
        'top': top if top is not None else sz,
        'left': sz,
        'bottom': bottom if bottom is not None else sz,
        'right': sz,
    }
    for side, s in sides.items():
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(s))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)


def set_cell_margin(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        mar.append(e)
    tcPr.append(mar)


def prevent_row_split(row):
    """행 가운데서 페이지 나뉘지 않도록."""
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


# ─── 상태 분류 ────────────────────────────────────────────
def classify_status(val: str):
    """
    반환: (symbol, weight)
      weight ∈ {'bold', 'normal', 'italic'}
    """
    v = val.strip()
    if not v:
        return '', 'normal'
    if v == '미출제':
        return '○', 'italic'
    if v.startswith('약출제'):
        return '◐', 'normal'
    return '●', 'bold'


# ─── 문서 준비 (가로 방향) ────────────────────────────────
doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width  = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = section.right_margin = Cm(1.8)
section.top_margin  = section.bottom_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
style.font.size = Pt(10)
style.font.color.rgb = BLACK


# ─── 타이틀 ───────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('정보처리기사 실기 출제현황 · 뜻 포함판')
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = BLACK
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2020~2025 기출 분석 · 2026 쪽집게 전략 포함')
run.font.size = Pt(11)
run.font.color.rgb = GRAY_MID
p.paragraph_format.space_after = Pt(8)


# ─── 범례 (흑백) ──────────────────────────────────────────
legend_items = [
    ('● 출제됨',  'bold',   '정답 칸에 직접 등장 (굵은 글씨)'),
    ('◐ 약출제', 'normal', '문제/보기/해설에서만 확인 (일반)'),
    ('○ 미출제', 'italic', '보유 기출 기준 미확인 (기울임)'),
]
leg = doc.add_table(rows=1, cols=3)
leg.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, weight, desc) in enumerate(legend_items):
    cell = leg.rows[0].cells[i]
    cell.width = Cm(8.7)
    set_cell_border(cell, '555555', 6)
    set_cell_margin(cell, 90, 90, 130, 130)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = cp.add_run(label)
    r1.font.size = Pt(11)
    r1.font.color.rgb = BLACK
    if weight == 'bold':
        r1.font.bold = True
    elif weight == 'italic':
        r1.font.italic = True
    r2 = cp.add_run(f'\n{desc}')
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = GRAY_MID
prevent_row_split(leg.rows[0])

doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ─── 헤딩 ────────────────────────────────────────────────
def add_h1(text):
    """대분류 (# / ##) — 이중 하단 보더로 강조, 페이지 나눔 없음."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    # 페이지 브레이크 금지: keep_with_next & 페이지 중간 허용
    pPr = p._p.get_or_add_pPr()
    keep = OxmlElement('w:keepNext')
    pPr.append(keep)
    # 하단 double border
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'double')
    bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = BLACK


def add_h2(text):
    """중분류 (### ) — 두꺼운 하단 single border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    keep = OxmlElement('w:keepNext')
    pPr.append(keep)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '555555')
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = BLACK


def add_h3(text):
    """소분류 추가 사용시."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = GRAY_DARK


def add_callout(lines, title=None):
    """인용구(>). 좌측 두꺼운 세로선 느낌의 회색 박스."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(26.1)
    set_cell_bg(cell, RGBColor(0xF5, 0xF5, 0xF5))
    # 왼쪽 굵은 검정 보더
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side, sz, col in (('top', 4, 'CCCCCC'), ('bottom', 4, 'CCCCCC'),
                          ('right', 4, 'CCCCCC'), ('left', 18, '000000')):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), col)
        borders.append(b)
    tcPr.append(borders)
    set_cell_margin(cell, 120, 120, 180, 180)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(0)
    if title:
        r = cp.add_run(title)
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = BLACK
        cp.add_run('\n')
    for i, line in enumerate(lines):
        if i > 0 or title:
            cp.add_run('\n')
        r = cp.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = GRAY_DARK
    prevent_row_split(tbl.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ─── 표 빌더 ──────────────────────────────────────────────
def make_table(headers, rows, status_col=None):
    """
    - 가용 폭 26.1cm 기준 프리셋.
    - status_col 지정 시: 상태 심볼(●◐○) 항목명에 prefix, 굵기·이탤릭 적용.
    - 컬러 사용 안 함. 출제됨 행은 살짝 진한 배경으로 구분.
    """
    n = len(headers)
    # 중요도 컬럼 포함 여부로 프리셋 선택
    has_importance = len(headers) >= 1 and headers[0].strip() == '중요도'
    width_presets = {
        2: [Cm(6.0),  Cm(20.1)],
        3: [Cm(4.8),  Cm(6.0),  Cm(15.3)],
        4: [Cm(2.3),  Cm(5.5),  Cm(5.5),  Cm(12.8)] if has_importance
           else [Cm(4.8),  Cm(6.0),  Cm(12.7), Cm(2.6)],
    }
    widths = width_presets.get(n, [Cm(26.1 / n)] * n)
    # 항목 컬럼 위치
    item_col = 0
    for idx, h in enumerate(headers):
        if h.strip() == '항목':
            item_col = idx
            break

    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    tblPr = tbl._tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)

    # ── 헤더 행 ──
    hdr = tbl.rows[0]
    hdr.height = Cm(0.75)
    prevent_row_split(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = widths[i]
        set_cell_bg(cell, HEADER_BG)
        set_cell_border(cell, 'FFFFFF', 6)
        set_cell_margin(cell, 80, 80, 110, 110)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = HEADER_TXT

    # ── 데이터 행 ──
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        prevent_row_split(tr)
        is_even = (ri % 2 == 1)

        # 상태 판단
        symbol, weight = '', 'normal'
        if status_col is not None and status_col < len(row):
            symbol, weight = classify_status(row[status_col])

        # 행 배경
        row_bg = None
        if weight == 'bold':
            row_bg = EMPHASIS_BG
        elif is_even:
            row_bg = ROW_ALT_BG

        for ci, text in enumerate(row):
            cell = tr.cells[ci]
            cell.width = widths[ci]
            set_cell_border(cell, 'BBBBBB', 4)
            set_cell_margin(cell, 55, 55, 110, 110)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_bg:
                set_cell_bg(cell, row_bg)

            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(0)

            # 중요도 컬럼 (has_importance일 때 ci==0): 가운데 정렬 + 굵게
            if has_importance and ci == 0:
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cp.add_run(text)
                r.font.size = Pt(10)
                r.font.color.rgb = BLACK
                r.font.bold = True
            # 항목 컬럼: 심볼 prefix + weight
            elif ci == item_col and symbol:
                rs = cp.add_run(f'{symbol} ')
                rs.font.size = Pt(10)
                rs.font.color.rgb = BLACK
                rs.font.bold = True
                r = cp.add_run(text)
                r.font.size = Pt(9.5)
                r.font.color.rgb = BLACK
                if weight == 'bold':
                    r.font.bold = True
                elif weight == 'italic':
                    r.font.italic = True
                    r.font.color.rgb = GRAY_DARK
            else:
                r = cp.add_run(text)
                r.font.size = Pt(9.5)
                r.font.color.rgb = BLACK
                if ci == status_col:
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if weight == 'bold':
                        r.font.bold = True
                    elif weight == 'italic':
                        r.font.italic = True
                        r.font.color.rgb = GRAY_MID
                elif ci == 0:
                    r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


# ─── 마크다운 파서 ─────────────────────────────────────────
def parse_md(text):
    lines = text.splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if stripped.startswith('# '):
            blocks.append(('h1', stripped[2:].strip()))
            i += 1
            continue
        if stripped.startswith('## '):
            blocks.append(('h2', stripped[3:].strip()))
            i += 1
            continue
        if stripped.startswith('### '):
            blocks.append(('h3', stripped[4:].strip()))
            i += 1
            continue
        if re.match(r'^-{3,}$', stripped):
            blocks.append(('hr',))
            i += 1
            continue
        if stripped.startswith('>'):
            qlines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                qlines.append(lines[i].strip().lstrip('>').strip())
                i += 1
            blocks.append(('quote', qlines))
            continue
        if stripped.startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|?\s*:?-+', lines[i + 1]):
            header = [c.strip() for c in stripped.strip('|').split('|')]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                if len(cells) < len(header):
                    cells += [''] * (len(header) - len(cells))
                elif len(cells) > len(header):
                    cells = cells[:len(header)]
                rows.append(cells)
                i += 1
            blocks.append(('table', header, rows))
            continue
        if re.match(r'^\d+\.\s', stripped):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                items.append(re.sub(r'^\d+\.\s', '', lines[i].strip()))
                i += 1
            blocks.append(('ol', items))
            continue
        if stripped.startswith('- '):
            items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                items.append(lines[i].strip()[2:])
                i += 1
            blocks.append(('ul', items))
            continue
        blocks.append(('para', stripped))
        i += 1

    return blocks


# ─── 인라인 처리 ─────────────────────────────────────────
def clean_inline(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


def render_inline_para(text, size=10, space_before=0, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    parts = re.split(r'(\*\*.+?\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.font.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Consolas'
        else:
            r = p.add_run(part)
        r.font.size = Pt(size)
        r.font.color.rgb = BLACK
    return p


def render_list(items, ordered=False):
    for idx, item in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        marker = f'{idx + 1}. ' if ordered else '• '
        r = p.add_run(marker)
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = BLACK
        parts = re.split(r'(\*\*.+?\*\*|`[^`]+`)', item)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2])
                r.font.bold = True
            elif part.startswith('`') and part.endswith('`'):
                r = p.add_run(part[1:-1])
                r.font.name = 'Consolas'
            else:
                r = p.add_run(part)
            r.font.size = Pt(10)
            r.font.color.rgb = BLACK


def detect_status_col(headers):
    for i, h in enumerate(headers):
        if h.strip() in ('상태', '출제 이력'):
            return i
    return None


# ─── 렌더링 ──────────────────────────────────────────────
md_text = SRC.read_text(encoding='utf-8')
blocks = parse_md(md_text)

skip_first_h1 = True

for b in blocks:
    kind = b[0]
    if kind == 'h1':
        if skip_first_h1:
            skip_first_h1 = False
            continue
        add_h1(clean_inline(b[1]))
    elif kind == 'h2':
        add_h1(clean_inline(b[1]))
    elif kind == 'h3':
        add_h2(clean_inline(b[1]))
    elif kind == 'hr':
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    elif kind == 'quote':
        add_callout([clean_inline(l) for l in b[1] if l])
    elif kind == 'para':
        render_inline_para(b[1])
    elif kind == 'ol':
        render_list(b[1], ordered=True)
    elif kind == 'ul':
        render_list(b[1], ordered=False)
    elif kind == 'table':
        headers = [clean_inline(h) for h in b[1]]
        rows = [[clean_inline(c) for c in row] for row in b[2]]
        status_col = detect_status_col(headers)
        make_table(headers, rows, status_col=status_col)


# ─── 저장 ─────────────────────────────────────────────────
doc.save(OUT)
print(f'✔ 저장 완료: {OUT}')
